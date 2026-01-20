"""
Ollama Service Module

Handles all interactions with the Ollama API for text generation,
document summarization, and other LLM tasks.
"""

import httpx
import logging
from typing import Optional, List, Dict, Any
import io

from app.config import settings

logger = logging.getLogger(__name__)


class OllamaService:
    """
    Service class for interacting with Ollama API
    
    Provides methods for text summarization, querying, and model management.
    """
    
    def __init__(self):
        self.base_url = settings.OLLAMA_URL
        self.default_model = settings.OLLAMA_DEFAULT_MODEL
        self.timeout = settings.OLLAMA_TIMEOUT
    
    async def check_health(self) -> bool:
        """
        Check if Ollama server is accessible
        
        Returns:
            True if Ollama is accessible, False otherwise
        """
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(
                    f"{self.base_url}/api/tags",
                    timeout=5.0
                )
                return response.status_code == 200
        except Exception as e:
            logger.warning(f"Ollama health check failed: {str(e)}")
            return False
    
    async def list_models(self) -> List[Dict[str, Any]]:
        """
        List all available models in Ollama
        
        Returns:
            List of model information dictionaries
        """
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(
                    f"{self.base_url}/api/tags",
                    timeout=10.0
                )
                response.raise_for_status()
                data = response.json()
                return data.get("models", [])
        except httpx.HTTPError as e:
            logger.error(f"Failed to list models: {str(e)}")
            raise Exception(f"Failed to list Ollama models: {str(e)}")
    
    async def generate(
        self,
        prompt: str,
        model: Optional[str] = None,
        system: Optional[str] = None,
        temperature: float = 0.7,
        max_tokens: Optional[int] = None,
        stream: bool = False
    ) -> str:
        """
        Generate text using Ollama
        
        Args:
            prompt: The prompt to generate from
            model: Model to use (defaults to configured default)
            system: Optional system prompt
            temperature: Generation temperature (0.0-1.0)
            max_tokens: Maximum tokens to generate
            stream: Whether to stream the response
            
        Returns:
            Generated text
        """
        model = model or self.default_model
        
        payload = {
            "model": model,
            "prompt": prompt,
            "stream": stream,
            "options": {
                "temperature": temperature
            }
        }
        
        if system:
            payload["system"] = system
        
        if max_tokens:
            payload["options"]["num_predict"] = max_tokens
        
        logger.info(f"Generating with model: {model}")
        
        try:
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    f"{self.base_url}/api/generate",
                    json=payload,
                    timeout=self.timeout
                )
                response.raise_for_status()
                data = response.json()
                return data.get("response", "")
        except httpx.HTTPError as e:
            logger.error(f"Ollama generation failed: {str(e)}")
            raise Exception(f"Ollama generation failed: {str(e)}")
    
    async def chat(
        self,
        messages: List[Dict[str, str]],
        model: Optional[str] = None,
        temperature: float = 0.7,
        stream: bool = False
    ) -> str:
        """
        Chat with Ollama using the chat API
        
        Args:
            messages: List of message dictionaries with 'role' and 'content'
            model: Model to use (defaults to configured default)
            temperature: Generation temperature (0.0-1.0)
            stream: Whether to stream the response
            
        Returns:
            Assistant's response
        """
        model = model or self.default_model
        
        payload = {
            "model": model,
            "messages": messages,
            "stream": stream,
            "options": {
                "temperature": temperature
            }
        }
        
        logger.info(f"Chat with model: {model}")
        
        try:
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    f"{self.base_url}/api/chat",
                    json=payload,
                    timeout=self.timeout
                )
                response.raise_for_status()
                data = response.json()
                return data.get("message", {}).get("content", "")
        except httpx.HTTPError as e:
            logger.error(f"Ollama chat failed: {str(e)}")
            raise Exception(f"Ollama chat failed: {str(e)}")
    
    async def summarize_text(
        self,
        text: str,
        model: Optional[str] = None,
        max_length: int = 500,
        style: str = "concise"
    ) -> str:
        """
        Summarize text using Ollama
        
        Args:
            text: Text to summarize
            model: Model to use (defaults to configured default)
            max_length: Approximate maximum length of summary in words
            style: Summary style ('concise', 'detailed', 'bullet_points', 'executive')
            
        Returns:
            Summary of the text
        """
        style_instructions = {
            "concise": "Provide a brief, concise summary that captures the main points.",
            "detailed": "Provide a comprehensive summary that covers all important details and nuances.",
            "bullet_points": "Provide the summary as a list of bullet points highlighting key information.",
            "executive": "Provide an executive summary suitable for business professionals, focusing on key findings, implications, and recommendations.",
            "academic": "Provide an academic summary with proper structure: background, methods (if applicable), key findings, and conclusions."
        }
        
        instruction = style_instructions.get(style, style_instructions["concise"])
        
        system_prompt = f"""You are a professional document summarizer. Your task is to create clear, accurate summaries.

Instructions:
- {instruction}
- Keep the summary to approximately {max_length} words
- Maintain the factual accuracy of the original text
- Do not add information not present in the original text
- Use clear, professional language
- Preserve important names, dates, and figures"""

        prompt = f"""Please summarize the following text:

---
{text}
---

Summary:"""

        logger.info(f"Summarizing text of length {len(text)} with style '{style}'")
        
        return await self.generate(
            prompt=prompt,
            model=model,
            system=system_prompt,
            temperature=0.3,  # Lower temperature for more consistent summaries
            max_tokens=max_length * 2  # Approximate tokens from words
        )
    
    async def query(
        self,
        prompt: str,
        model: Optional[str] = None,
        context: Optional[str] = None
    ) -> str:
        """
        Send a custom query to Ollama
        
        Args:
            prompt: The query/prompt
            model: Model to use (defaults to configured default)
            context: Optional context to include
            
        Returns:
            Model response
        """
        full_prompt = prompt
        if context:
            full_prompt = f"""Context:
{context}

Question/Task:
{prompt}

Response:"""
        
        return await self.generate(
            prompt=full_prompt,
            model=model,
            temperature=0.7
        )
    
    async def extract_pdf_text(self, content: bytes) -> str:
        """
        Extract text from PDF content
        
        Args:
            content: PDF file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            import pypdf
            
            pdf_file = io.BytesIO(content)
            reader = pypdf.PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return "\n\n".join(text_parts)
        except ImportError:
            logger.error("pypdf not installed. Install with: pip install pypdf")
            raise Exception("PDF processing not available. pypdf library required.")
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    async def extract_docx_text(self, content: bytes) -> str:
        """
        Extract text from DOCX content
        
        Args:
            content: DOCX file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise Exception("DOCX processing not available. python-docx library required.")
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    async def analyze_document(
        self,
        text: str,
        analysis_type: str = "general",
        model: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Analyze a document for various purposes
        
        Args:
            text: Document text to analyze
            analysis_type: Type of analysis ('general', 'sentiment', 'entities', 'topics')
            model: Model to use
            
        Returns:
            Analysis results
        """
        analysis_prompts = {
            "general": """Analyze the following document and provide:
1. Main topic/subject
2. Key points (3-5 points)
3. Target audience
4. Tone and style
5. Brief summary (2-3 sentences)""",
            
            "sentiment": """Analyze the sentiment of the following document. Provide:
1. Overall sentiment (positive/negative/neutral/mixed)
2. Sentiment score (-1.0 to 1.0)
3. Key positive aspects
4. Key negative aspects
5. Emotional tone""",
            
            "entities": """Extract key entities from the following document:
1. People mentioned
2. Organizations
3. Locations
4. Dates/Times
5. Important numbers/statistics""",
            
            "topics": """Identify the topics discussed in the following document:
1. Primary topic
2. Secondary topics
3. Keywords
4. Subject categories
5. Related themes"""
        }
        
        prompt_template = analysis_prompts.get(analysis_type, analysis_prompts["general"])
        
        prompt = f"""{prompt_template}

Document:
---
{text}
---

Analysis:"""

        response = await self.generate(
            prompt=prompt,
            model=model,
            temperature=0.3
        )
        
        return {
            "analysis_type": analysis_type,
            "result": response
        }
